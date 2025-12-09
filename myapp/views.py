import random
import re
import json
import time
import tempfile
import os
import docx

from .models import Question, Requirement, QuestionAnswer, User, Candidate
from .serializers import (
    AnswerSerializer, 
    QuestionSerializer, 
    HrSerializer, 
    PhotoSerializer, 
    RequirementSerializer, 
    PublicCandidateSerializer, 
    RegisterSerializer, 
    LoginSerializer,
    ChatAiSerializer
)
from .utils import analyze_facial_expressions
from .tasks import rate_answer

from rest_framework import status
from rest_framework.response import Response
from rest_framework.views import APIView
from PyPDF2 import PdfReader
from django.conf import settings
from django.core.mail import send_mail
from django.core.exceptions import ValidationError
from io import BytesIO
import google.generativeai as genai
from django.contrib.auth import authenticate, login as auth_login
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework.permissions import IsAuthenticated, AllowAny


genai.configure(api_key=os.getenv("GEMINI_API_KEY"))


class RegisterView(APIView):
    serializer_class = RegisterSerializer

    def get(self, request):
        users = User.objects.all()
        serializer = RegisterSerializer(users, many=True)
        return Response(serializer.data)
    
    def post(self, request):
        data = request.data
        try:
            User.objects.get(username=data['username'])
            return Response({"status": "error", "message": f"User with this username address already exists."}, status=400)
        except:
            serializer = self.serializer_class(data=data)
            if serializer.is_valid():
                serializer.save()
                return Response({"status": "success", "message": "User created successfully!","data":serializer.data}, status=200)
            
            return Response(serializer.errors, status=400)
        
    def put(self, request, pk=None):
        data = request.data
        try:
            user_obj = User.objects.get(id=pk)
        except User.DoesNotExist:
            return Response({"status": "error", "message": f"User Does not exist with ID: {pk}"}, status=400)
        
        serializer = RegisterSerializer(user_obj, data=data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response({"status": "success", "message": "Profile updated successfully!"}, status=200)
        return Response(serializer.errors, status=400)

    
class LoginView(APIView):
    serializer_class = LoginSerializer
    
    def post(self, request):
        data = request.data
        serializer = self.serializer_class(data=data)
        if serializer.is_valid():
            email = serializer.validated_data.get('email')
            password = serializer.validated_data.get('password')
            
            try:
                user_obj = User.objects.get(email=email)
                if not user_obj.is_active:
                    return Response({"status": "error", "message": "Your account has been disabled!"}, status=status.HTTP_400_BAD_REQUEST)
            except User.DoesNotExist:
                return Response({"status": "error", "message": "The email provided is invalid."}, status=status.HTTP_400_BAD_REQUEST)
            
            user = authenticate(email=email, password=password)
            if user is not None:
                auth_login(request, user)
                refresh = RefreshToken.for_user(user)
                
                response = {
                    'refresh': str(refresh),
                    'access': str(refresh.access_token),
                    'user': {
                        'id': user.id,
                        'username': user.username,
                        'email': user.email,
                        'role': user.role,
                        'is_staff': user.is_staff,
                        }
                }
                return Response(response, status=status.HTTP_200_OK)
            return Response({"status": "error", "message": "The password provided is invalid."}, status=status.HTTP_400_BAD_REQUEST)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class QuestionListAPIView(APIView):
    def get(self, request, pk=None):
        print("request==========", request.GET.get('candidate'))
        if pk:
            try:
                print("================", pk)
                candidate = Candidate.objects.get(pk=pk)
            except Candidate.DoesNotExist:
                return Response({"error": "Candidate not found."}, status=status.HTTP_404_NOT_FOUND)
            question = Question.objects.filter(candidate=candidate)
            if question:
                serializer = QuestionSerializer(question, many=True)
                return Response(serializer.data)
            else:
                questions = list(Question.objects.all().exclude(candidate=candidate))
                if not questions:
                    return Response({"message": "No questions available."}, status=404)
                random_questions = random.sample(questions, min(len(questions), 10))

                serializer = QuestionSerializer(random_questions, many=True)
                return Response(serializer.data)
        return Response({"error": "Candidate not found."}, status=status.HTTP_404_NOT_FOUND)

    def put(self, request, pk=None):
        """Update a question's text and technology"""
        if not pk:
            return Response({"error": "Question ID is required for update"}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            question = Question.objects.get(pk=pk, is_deleted=False)
        except Question.DoesNotExist:
            return Response({"error": "Question not found."}, status=status.HTTP_404_NOT_FOUND)
        
        # Get the fields to update
        question_text = request.data.get('text')
        technology = request.data.get('technology')
        # difficulty_level = request.data.get('difficulty_level')
        time_limit = request.data.get('time_limit')
        
        # Update fields if provided
        if question_text is not None:
            question.text = question_text
        if technology is not None:
            # Validate technology against choices
            valid_technologies = [choice[0] for choice in TECHNOLOGY_CHOICES]
            if technology not in valid_technologies:
                return Response({"error": f"Invalid technology. Must be one of: {valid_technologies}"}, status=status.HTTP_400_BAD_REQUEST)
            question.technology = technology
        if difficulty_level is not None:
            # Validate difficulty_level against choices
            valid_difficulties = [choice[0] for choice in DIFFICULTY_CHOICES]
            if difficulty_level not in valid_difficulties:
                return Response({"error": f"Invalid difficulty level. Must be one of: {valid_difficulties}"}, status=status.HTTP_400_BAD_REQUEST)
            question.difficulty_level = difficulty_level
        if time_limit is not None:
            try:
                question.time_limit = int(time_limit)
            except (ValueError, TypeError):
                return Response({"error": "time_limit must be a valid integer"}, status=status.HTTP_400_BAD_REQUEST)
        
        question.save()
        serializer = QuestionSerializer(question)
        return Response(serializer.data, status=status.HTTP_200_OK)

    def post(self, request):
        file = request.FILES.get('file')
        
        if file:
            try:
                filename = file.name.lower()
                questions = []

                if filename.endswith('.txt'):
                    file_content = file.read().decode('utf-8', errors='ignore')
                    questions = [q.strip() for q in file_content.split('\n') if q.strip()]

                elif filename.endswith('.pdf'):
                    pdf_reader = PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    questions = re.split(r'\n\s*\d+[\).]|\n*Q\d+[\).]?', text)
                    questions = [q.strip() for q in questions if len(q.strip()) > 5]
                else:
                    return Response({"error": "Unsupported file type. Please upload .txt or .pdf"}, status=status.HTTP_400_BAD_REQUEST)

                saved_questions = []
                for question_text in questions:
                    question_data = {
                        'text': question_text,
                        'hr': request.data.get('hr')
                    }

                    serializer = QuestionSerializer(data=question_data)
                    if serializer.is_valid():
                        serializer.save()
                        saved_questions.append(serializer.data)

                if not saved_questions:
                    return Response({"error": "No valid questions found in the file."}, status=status.HTTP_400_BAD_REQUEST)

                return Response(saved_questions, status=status.HTTP_201_CREATED)

            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        serializer = QuestionSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            
            return Response(serializer.data, status=status.HTTP_201_CREATED)
            
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class AnswerSaveView(APIView):
    def post(self, request):
        serializer = AnswerSerializer(data=request.data)
        if serializer.is_valid():
            answer = serializer.save()

            # Queue background task
            rate_answer.delay(answer.id)  

            return Response(
                {"saved_data": serializer.data, "message": "Answer saved, rating in progress..."},
                status=status.HTTP_201_CREATED
            )

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class CandidateView(APIView):
    def get_permissions(self):
        # Allow anyone to access GET; require auth for other methods
        if self.request.method == 'GET':
            return [AllowAny()]
        elif self.request.method == 'PUT':
            return [AllowAny()]
        return [IsAuthenticated()]
    
    def get(self,request, pk=None):
        if pk:
            try:
                hr_objects = Candidate.objects.get(pk=pk)
            except (ValueError, ValidationError):
                return Response({"error": "Invalid id format. Must be a UUID."}, status=status.HTTP_400_BAD_REQUEST)
            except Candidate.DoesNotExist:
                return Response({"error": "HR record not found."}, status=status.HTTP_404_NOT_FOUND)
            serializer = PublicCandidateSerializer(hr_objects)
            return Response(serializer.data, status=status.HTTP_200_OK)
        else:
            hr_objects = Candidate.objects.all().order_by('-created_at')
            serializer = HrSerializer(hr_objects, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)

    def post(self, request):
        files = request.FILES.getlist('upload_doc')
        if not files:
            return Response({"error": "No file uploaded"}, status=status.HTTP_400_BAD_REQUEST)

        created_records = []

        for file in files:
            text = ""
            try:
                filename = file.name.lower()
                if filename.endswith('.pdf'):
                    reader = PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                elif filename.endswith('.docx'):
                    document = docx.Document(file)
                    for para in document.paragraphs:
                        if para.text.strip():
                            text += para.text + "\n"
                else:
                    return Response({"error": "Unsupported file type. Please upload .pdf or .docx"}, status=status.HTTP_400_BAD_REQUEST)
            except Exception as e:
                return Response({"error": f"Failed to read resume: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST)

            if not text.strip():
                return Response({"error": "Could not extract any readable text from the uploaded document."}, status=status.HTTP_400_BAD_REQUEST)

            prompt = f"""
                    Extract the following fields from this resume text: 
                    - name
                    - email
                    - phone
                    - technology (take only one main technology)
                    - experience (total experience)
                    - companies (array of objects with company_name, start_date, end_date)
                    
                    Date Requirements:
                    - Convert ALL dates to YYYY-MM format (e.g., 2024-07, 2025-03)
                    - If only year is available, use YYYY format (e.g., 2024)
                    - If date is unclear, use null
                    - Handle various formats: "July 2024", "12/03/2022", "2022-2024", etc.
                    - IMPORTANT: If candidate is currently working at company (end date shows "present", "current", "till date", "ongoing", etc.), set end_date = "running"
                    
                    Return only valid JSON (no explanations, no ```json blocks) with this format:
                    {{
                        "name": "John Doe",
                        "email": "john@example.com",
                        "phone": "+1234567890",
                        "technology": "React",
                        "experience": "2 years 3 months",
                        "companies": [
                            {{
                                "company_name": "Tech Corp",
                                "start_date": "2022-01",
                                "end_date": "2024-03"
                            }},
                            {{
                                "company_name": "Current Company Inc",
                                "start_date": "2024-04",
                                "end_date": "running"
                            }}
                        ]
                    }}
                    
                    Resume text:
                    {text}
                """
            print(prompt)

            # Try LLM with simple retries
            result_text = None
            last_error = None
            for attempt in range(3):
                try:
                    # Create the model instance
                    model = genai.GenerativeModel("models/gemini-2.5-flash")
                    # Generate content
                    response = model.generate_content(
                        prompt,
                        generation_config=genai.types.GenerationConfig(
                            temperature=0.1,
                            response_mime_type="application/json",  # or "text/plain" if you want plain text
                        ),
                    )
                    # Extract text safely
                    result_text = getattr(response, "text", None)
                    print("===================")
                    # print(result_text)
                    if result_text:
                        result_text = result_text.strip()
                    break  # exit loop on success

                except Exception as e:
                    last_error = str(e)
                    # If it's a quota error, don't retry
                    if "429" in last_error or "quota" in last_error.lower():
                        break
                    print(f"Attempt {attempt+1} failed: {e}")
                    time.sleep(1.0)  # backoff before retrying

            data = None
            if result_text:
                try:
                    clean_text = re.sub(r"```json|```", "", result_text).strip()
                    print(clean_text)
                    data = json.loads(clean_text)
                except Exception as e:
                    # fall back to heuristic parsing if LLM output is malformed
                    last_error = e

            fallback_used = False
            if not data:
                fallback_used = True
                lowered = text.lower()
                email_match = re.search(r"[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,}", text)
                phone_match = re.search(r"(\+?\d[\d\s\-\(\)]{7,}\d)", text)
                first_line = next((ln.strip() for ln in text.splitlines() if ln and len(ln.strip()) > 1), "")
                tech_map = ["python", ".net", "java", "react"]
                tech = next((t for t in tech_map if t in lowered), "")
                
                # Extract companies using regex in fallback
                companies = []
                # Look for company names patterns
                company_patterns = [
                    r"(?:worked at|employed by|experience at|company|private|limited|ltd|inc|corp|technologies|solutions|systems|services)\s+([A-Za-z0-9\s&]+?)(?:\n|,|\.|\s+from|\s+since|\s+to|\s+till|\s+-|\s*\(|\s*[0-9]{4})",
                    r"([A-Za-z0-9\s&]+?(?:technologies|solutions|systems|services|limited|ltd|inc|corp))\s*(?:\n|,|\.|\s+from|\s+since|\s+to|\s+till|\s+-|\s*\(|\s*[0-9]{4})",
                    r"^\s*([A-Za-z0-9\s&]{2,50})\s*\n.*?(?:[0-9]{4}|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)"
                ]
                
                for pattern in company_patterns:
                    matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
                    for match in matches:
                        company_name = match.strip() if isinstance(match, str) else match[0].strip() if match else ""
                        if len(company_name) > 2 and len(company_name) < 100:
                            # Avoid duplicates
                            if not any(comp.get("company_name", "").lower() == company_name.lower() for comp in companies):
                                companies.append({
                                    "company_name": company_name,
                                    "start_date": None,
                                    "end_date": None
                                })
                
                # Limit to first 5 companies to avoid too much data
                companies = companies[:5]
                
                data = {
                    "name": first_line[:100],
                    "email": email_match.group(0) if email_match else "",
                    "phone": phone_match.group(0) if phone_match else "",
                    "technology": tech,
                    "companies": companies,
                }
                if not any([data.get("name"), data.get("email"), data.get("phone"), data.get("technology")]):
                    # If even fallback got nothing useful, respond gracefully with 503 if provider failed, else 422
                    status_code = status.HTTP_503_SERVICE_UNAVAILABLE if last_error else status.HTTP_422_UNPROCESSABLE_ENTITY
                    return Response(
                        {
                            "error": "Could not extract fields from resume.",
                            "details": str(last_error) if last_error else "",
                        },
                        status=status_code,
                    )
            email = data.get("email")
            if email and Candidate.objects.filter(email=email).exists():
                return Response(
                    {
                        "error": "Email already exists.",
                        "details": "Email already exists.",
                    },
                    status=400,
                )

            # Store raw extracted document text in base_text
            llm_response = data if not fallback_used else {"fallback": True, "data": data}
            
            # Debug logging
            print(f"Extracted companies: {data.get('companies', [])}")
            print(f"Number of companies: {len(data.get('companies', []))}")
            
            hr_obj = Candidate.objects.create(
                upload_doc=file,
                name=data.get("name", ""),
                email=data.get("email", ""),
                phone=data.get("phone", ""),
                experience=data.get("experience", ""),
                technology=data.get("technology", ""),
                company=data.get("companies", []),
                base_text=text,  # Store raw document text instead of LLM response
            )
            frontend_base = "http://localhost:5173"
            hr_obj.shine_link = f"{frontend_base}/{hr_obj.id}/"
            hr_obj.save(update_fields=['shine_link'])
            created_records.append(hr_obj)


        serializer = HrSerializer(created_records, many=True)
        return Response({"records": serializer.data}, status=status.HTTP_201_CREATED)

    def put(self, request, pk):
        try:
            hr_obj = Candidate.objects.get(pk=pk)
        except Candidate.DoesNotExist:
            return Response({"error": "Candidate not found"}, status=status.HTTP_404_NOT_FOUND)

        serializer = HrSerializer(hr_obj, data=request.data, partial=True)
        if serializer.is_valid():
            interview_time = serializer.validated_data.get("time")
            interview_status = serializer.validated_data.get("interview_status")
            interview_quick = serializer.validated_data.get("is_quick")
            requirement = serializer.validated_data.get("requirement")
            is_selected = serializer.validated_data.get("is_selected")

            # If time is added, mark as scheduled and send email
            if interview_time:
                hr_obj.interview_status = "Scheduled"
                hr_obj.save(update_fields=["interview_status"])

                # Save serializer data (to ensure gmeet_link etc. are saved)
                serializer.save()

                # Send email with interview link
                subject = f"Interview Scheduled - {hr_obj.technology} Role"
                message = (
                    f"Hello {hr_obj.name},\n\n"
                    f"Your interview for the {hr_obj.technology} position has been scheduled.\n\n"
                    f"🗓 Date & Time: {interview_time.strftime('%A, %d %B %Y at %I:%M %p')}\n"
                    f"🔗 Google Meet Link: {hr_obj.gmeet_link or 'Not provided'}\n"
                    f"🔗 Share Link: {hr_obj.shine_link or 'Not provided'}\n"
                    f"Please make sure to join the interview on time.\n\n"
                    f"Best regards,\n"
                    f"HR Team"
                )

                if hr_obj.email:
                    try:
                        send_mail(
                            subject=subject,
                            message=message,
                            from_email=settings.DEFAULT_FROM_EMAIL,
                            recipient_list=[hr_obj.email],
                            fail_silently=False,
                        )
                    except Exception as e:
                        print(f"Email sending failed: {e}")

                return Response(serializer.data, status=status.HTTP_200_OK)
            if interview_status == "Completed":
                print("hr_obj:", hr_obj)
                print("hr_obj.id:", hr_obj.id)
                answers = list(
                    QuestionAnswer.objects.filter(candidate_id=hr_obj.id)
                    .values_list("answer_text", flat=True)
                )
                print("Answer------=-", answers)
                if not answers:
                    return Response({"message": "No answers found"}, status=200)

                prompt = (
                    "Evaluate the candidate's communication skills based on the provided answers.\n"
                    "Return ONLY a valid JSON object named communication_point containing:\n"
                    "- Grammar: integer score (0-10)\n"
                    "- ProfessionalLanguage: integer score (0-10)\n"
                    "- OverallGrammarExplanation: short paragraph explaining the grammar score.\n"
                    "- OverallProfessionalLanguageExplanation: short paragraph explaining the professional language score.\n"
                    "- OverallLanguageUsed: describe which language(s) the candidate used (e.g., English, Hindi, mixed).\n\n"
                    "Example format:\n"
                    "{\n"
                    '  "communication_point": {\n'
                    '    "Grammar": 8,\n'
                    '    "ProfessionalLanguage": 7,\n'
                    '    "OverallGrammarExplanation": "Grammar was mostly correct, with few minor sentence structure issues.",\n'
                    '    "OverallProfessionalLanguageExplanation": "The candidate used formal language but with occasional informal phrases.",\n'
                    '    "OverallLanguageUsed": "English"\n'
                    "  }\n"
                    "}\n\n"
                    f"Answers: {json.dumps(answers)}"
                )

                try:
                    model = genai.GenerativeModel("models/gemini-2.5-flash-lite")

                    # Generate JSON response
                    response = model.generate_content(
                        prompt,
                        generation_config=genai.types.GenerationConfig(
                            temperature=0.1,
                            response_mime_type="application/json",
                        ),
                    )
                    result_text = response.text.strip()
                except Exception as e:
                    print("error", e)
                    result_text = None

                print("Result Text", result_text)
                if not result_text:
                    return Response({"error": "Failed to generate communication evaluation"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                result_json = json.loads(result_text)
                if "communication_point" in result_json:
                    result_json = result_json["communication_point"]
                    hr_obj.communication = result_json
                    hr_obj.save(update_fields=["communication"])
                
                facial = analyze_facial_expressions(hr_obj)

            if interview_quick and str(hr_obj.requirement) != str(requirement):
                try:
                    prompt = f"""
                    You are an expert technical interviewer.

                    Generate 10 short, clear, challenging interview questions based on:

                    ### Requirement / Job Description
                    {requirement.base_text}

                    ### Candidate Information
                    Name: {hr_obj.name}
                    Email: {hr_obj.email}
                    Experience: {hr_obj.experience} years
                    Technology: {hr_obj.technology}

                    ### Rules:
                    - Focus questions on required technologies, tools, frameworks & experience level
                    - Include mix of theoretical + practical scenario questions
                    - Make questions specific, not generic
                    - Avoid yes/no questions
                    - Do NOT repeat similar questions
                    - Keep each question under 20 words

                    ### Output JSON Format ONLY:
                    {{
                    "questions": [
                        "Question 1 tell me about yourself",
                        "Question 2",
                        ...
                        "Question 10"
                    ]
                    }}
                    """
                    model = genai.GenerativeModel("models/gemini-2.5-flash-lite")
                    response = model.generate_content(
                        prompt,
                        generation_config=genai.types.GenerationConfig(
                            temperature=0.1,
                            response_mime_type="application/json",
                        ),
                    )

                    result = json.loads(response.text)
                    questions = result.get("questions", [])

                    question_objs = [Question(candidate=hr_obj, text=q) for q in questions ]

                    Question.objects.bulk_create(question_objs)
                except Exception as e:
                    print(f"Failed to generate questions: {e}")
            if is_selected == True or is_selected == False:
                hr_obj.save(update_fields=["is_selected"])
                serializer.save()

                subject = f"Interview Result"
                if is_selected == True:
                    message = (
                        f"Hello {hr_obj.name},\n\n"
                        f"Congratulations! You're selected in the interview.\n"
                        f"Hr is connected you soon.\n"
                        f"Best regards,\n"
                        f"HR Team."
                    )
                elif is_selected == False:
                    message = (
                        f"Hello {hr_obj.name},\n\n" 
                        f"Unfortunately, you are not selected in the interview.\n" 
                        f"Good luck for next interview.\n" 
                        f"Best regards,\n" 
                        f"HR Team."
                    )
                if hr_obj.email:
                    try:
                        send_mail(
                            subject=subject,
                            message=message,
                            from_email=settings.DEFAULT_FROM_EMAIL,
                            recipient_list=[hr_obj.email],
                            fail_silently=False,
                        )
                    except Exception as e:
                        print(f"Email sending failed: {e}")

                return Response(serializer.data, status=status.HTTP_200_OK)
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class PhotoView(APIView):
    def get(self, request, pk):
        try:
            hr_obj = Candidate.objects.get(pk=pk)
        except Candidate.DoesNotExist:
            return Response({"error": "Candidate not found"}, status=status.HTTP_404_NOT_FOUND)
        serializer = PhotoSerializer(hr_obj.photos.all(), many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)

    def post(self, request, pk):
        try:
            hr_obj = Candidate.objects.get(pk=pk)
        except Candidate.DoesNotExist:
            return Response({"error": "Candidate not found"}, status=status.HTTP_404_NOT_FOUND)
        serializer = PhotoSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            hr_obj.photos.add(serializer.instance)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class RequirementView(APIView):
    permission_classes = [IsAuthenticated]
    def get(self, request):
        requirement_objects = Requirement.objects.filter(is_deleted=False).order_by('-created_at')
        serializer = RequirementSerializer(requirement_objects, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)

    def post(self, request):
        file = request.FILES.get('file')
        if not file:
            return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)
            
        try:
            file_data = file.read()
            filename = file.name
            file_extension = filename.split('.')[-1].lower()

            text = ""

            # Handle DOCX files
            if file_extension == 'docx':
                docx_file = BytesIO(file_data)
                doc = docx.Document(docx_file)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"

            # Handle DOC files (requires textract)
            elif file_extension == 'doc':
                with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
                    temp_file.write(file_data)
                    temp_file_path = temp_file.name

                try:
                    text = textract.process(temp_file_path).decode('utf-8')
                finally:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)
            # Handle PDF files
            elif file_extension == 'pdf':
                try:
                    pdf_reader = PdfReader(BytesIO(file_data))
                    for page in pdf_reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
                except Exception as e:
                    return Response({"error": f"Failed to read PDF: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST)
            else:
                return Response({"error": "Unsupported file type. Please upload .docx, .doc, or .pdf"}, status=status.HTTP_400_BAD_REQUEST)
            if not text.strip():
                return Response({"error": "Could not extract any text from the uploaded file"}, 
                             status=status.HTTP_400_BAD_REQUEST)

            # Prepare the prompt for OpenAI
            prompt = """
            Extract the following fields from the given job requirement text:
            - name: Job title/requirement name
            - experience: Required experience (e.g., "2-5 years")
            - technology: Main technology/stack required (e.g., "Python", "React")
            - No_of_openings: Number of open positions (extract as integer)
            - notice_period: Notice period in days (extract as integer)
            - priority: Boolean indicating if this is a high priority requirement (true/false)

            Return only a valid JSON object with these fields. If a field cannot be determined, use null.
            Example output:
            {
                "name": "Senior Python Developer",
                "experience": "2-5 years",
                "technology": "Python",
                "No_of_openings": 3,    
                "notice_period": 30,
                "priority": true
            }

            Here is the requirement text to analyze:
            """ + text

            # Call OpenAI API
            model = genai.GenerativeModel("models/gemini-2.5-flash")
            # Generate JSON response
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.1,
                    response_mime_type="application/json",
                ),
            )
            result_text = response.text.strip()
            
            # Parse the response
            clean_text = re.sub(r'```json|```', '', result_text).strip()
            data = json.loads(clean_text)

            if not any([data.get("name"), data.get("experience"), data.get("technology"), data.get("No_of_openings"), data.get("notice_period"), data.get("priority")]):
                # If even fallback got nothing useful, respond gracefully with 503 if provider failed, else 422
                status_code = status.HTTP_503_SERVICE_UNAVAILABLE if last_error else status.HTTP_422_UNPROCESSABLE_ENTITY
                return Response(
                    {
                        "error": "Could not extract fields from Document.",
                        "details": str(last_error) if last_error else "",
                    },
                    status=status_code,
                )
            
            # Create the requirement with extracted data
            requirement = Requirement.objects.create(
                file=file,
                base_text=text,
                name=data.get('name'),
                experience=data.get('experience'),
                technology=data.get('technology'),
                No_of_openings=data.get('No_of_openings'),
                notice_period=data.get('notice_period'),
                priority=bool(data.get('priority', False))
            )
            
            serializer = RequirementSerializer(requirement)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
            
        except json.JSONDecodeError as e:
            return Response(
                {"error": f"Failed to parse AI response: {str(e)}", "raw_response": result_text},
                status=status.HTTP_422_UNPROCESSABLE_ENTITY
            )
        except Exception as e:
            return Response(
                {"error": f"Failed to process requirement: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def put(self, request, pk):
        try:
            requirement = Requirement.objects.get(pk=pk)
        except Requirement.DoesNotExist:
            return Response(
                {"error": "Requirement not found"},
                status=status.HTTP_404_NOT_FOUND
            )

        serializer = RequirementSerializer(requirement, data=request.data)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk=None):
        """Soft delete a requirement by setting is_deleted=True"""
        try:
            requirement = Requirement.objects.get(pk=pk)
        except (ValueError, ValidationError):
            return Response({"error": "Invalid id format. Must be a UUID."}, status=status.HTTP_400_BAD_REQUEST)
        except Requirement.DoesNotExist:
            return Response({"error": "Requirement not found."}, status=status.HTTP_404_NOT_FOUND)

        requirement.is_deleted = True
        requirement.save(update_fields=['is_deleted'])
        return Response(status=status.HTTP_204_NO_CONTENT)


class JDAssistantView(APIView):
    permission_classes = [IsAuthenticated]
    
    def post(self, request, action):
        """
        Handle different JD assistant actions:
        - analyze: Extract and validate JD fields from user message
        - generate: Generate complete JD based on fields
        - save: Save generated JD to database
        """
        if action == 'analyze':
            return self.analyze_input(request)
        elif action == 'generate':
            return self.generate_jd(request)
        elif action == 'save':
            return self.save_jd(request)
        else:
            return Response({"error": "Invalid action"}, status=status.HTTP_400_BAD_REQUEST)
    
    def analyze_input(self, request):
        """Analyze user message and extract JD fields"""
        message = request.data.get('message', '')
        if not message:
            return Response({"error": "Message is required"}, status=status.HTTP_400_BAD_REQUEST)
        
        prompt = f"""
        Analyze this job description request and extract structured information.
        
        Required fields to extract:
        - name (Job title/position name)  -- this is OPTIONAL, you may infer it from technology/experience
        - experience (Required experience in years or range)
        - technology (Primary technology stack)
        - No_of_openings (Number of positions, integer)
        - notice_period (Notice period in days, integer)
        - priority (Boolean: true if urgent/high priority)
        
        User message: "{message}"
        
        Return only valid JSON with this format:
        {{
            "status": "ready" or "need_more_info",
            "fields": {{
                "name": "extracted JD name or null",
                "experience": "extracted value or null", 
                "technology": "extracted value or null",
                "No_of_openings": extracted_integer_or_null,
                "notice_period": extracted_integer_or_null,
                "priority": true_or_false_or_null
            }},
            "missing_fields": ["list", "of", "missing", "field", "names"]
        }}
        
        Rules:
        - REQUIRED fields are only: experience and technology
        - The JD name (field "name") is OPTIONAL. If possible, infer a good JD name, otherwise leave it null.
        - Set status to "ready" only if experience AND technology are provided
        - Set status to "need_more_info" if either experience or technology is missing
        - For No_of_openings, notice_period: extract numbers or set null
        - For priority: detect words like "urgent", "immediate", "high priority" as true
        - Include ONLY the actually missing required fields (experience, technology) in missing_fields array
        """
        
        try:
            model = genai.GenerativeModel("models/gemini-2.5-flash")
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.1,
                    response_mime_type="application/json",
                ),
            )
            
            result_text = getattr(response, "text", None)
            if result_text:
                result_text = result_text.strip()
                clean_text = re.sub(r"```json|```", "", result_text).strip()
                data = json.loads(clean_text)
                return Response(data, status=status.HTTP_200_OK)
            else:
                return Response({"error": "Failed to analyze input"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            return Response({"error": f"Analysis failed: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    def generate_jd(self, request):
        """Generate complete job description based on provided fields"""
        # Accept either `fields` (preferred) or `analysis_data` (from frontend)
        raw_fields = request.data.get('fields')
        analysis_data = request.data.get('analysis_data')

        # If analysis_data is present, support both:
        # 1) { status, fields: {...}, missing_fields }
        # 2) { name, experience, technology, No_of_openings, ... } (flat shape from frontend)
        if not raw_fields and isinstance(analysis_data, dict):
            inner_fields = analysis_data.get('fields')
            if isinstance(inner_fields, dict):
                raw_fields = inner_fields
            else:
                # Fallback: treat the full analysis_data as the fields dict
                raw_fields = analysis_data

        fields = raw_fields or {}

        # Only experience and technology are required; name is optional and can be auto-suggested
        name = fields.get('name')
        experience = fields.get('experience')
        technology = fields.get('technology')

        if not experience or not technology:
            return Response(
                {
                    "error": "Missing required fields",
                    "details": "Fields 'experience' and 'technology' are required to generate JD.",
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Auto-suggest JD name if not provided
        if not name and technology:
            name = f"{technology} Developer"
            fields["name"] = name

        prompt = f"""
        Generate a comprehensive job description based on these details:

        Job Title: {name}
        Experience Required: {experience}
        Technology: {technology}
        Number of Openings: {fields.get('No_of_openings', 'Not specified')}
        Notice Period: {fields.get('notice_period', 'Not specified')} days
        Priority: {'High Priority' if fields.get('priority') else 'Normal'}

        Generate a complete job description with:
        1. Job Summary
        2. Responsibilities  
        3. Requirements (technical and soft skills)
        4. Nice-to-have skills
        5. What we offer

        Format as clean text without markdown formatting.
        """

        try:
            model = genai.GenerativeModel("models/gemini-2.5-flash")
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.3,
                    response_mime_type="text/plain",
                ),
            )

            jd_text = getattr(response, "text", None)
            if jd_text:
                return Response(
                    {
                        "fields": fields,
                        "jd_text": jd_text.strip(),
                    },
                    status=status.HTTP_200_OK,
                )
            else:
                return Response(
                    {"error": "Failed to generate job description"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                )

        except Exception as e:
            return Response(
                {"error": f"Generation failed: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )
    
    def save_jd(self, request):
        """Save generated job description to Requirement model"""
        fields = request.data.get('fields', {})
        jd_text = request.data.get('jd_text', '')

        experience = fields.get('experience')
        technology = fields.get('technology')

        # If required fields are missing, try to extract them from jd_text using Gemini
        if (not experience or not technology) and jd_text:
            # Truncate very long JDs before sending to LLM to reduce latency
            # Usually the top part of the JD is enough to infer experience/technology
            jd_snippet = jd_text[:2000]

            extract_prompt = f"""
            From the following job description text, extract structured fields.

            VERY IMPORTANT INSTRUCTIONS FOR THE "technology" FIELD:
            - Return ONLY the 1 to 3 MAIN technologies / stacks.
            - Each technology name must be short (for example: "Python", "AWS", "React").
            - Do NOT include versions, frameworks in brackets, long tool lists or full stacks.
            - Do NOT return more than 3 items.
            - If there are many skills, pick only the top 2-3 most central technologies.
            - Join them in a single short comma-separated string, for example: "Python, AWS".

            JD Text (may be truncated):
            {jd_snippet}

            Return ONLY valid JSON in this format:
            {{
                "name": "Job title or null",
                "experience": "Experience requirement or null",
                "technology": "1-3 short main technologies as a comma-separated string or null",
                "No_of_openings": integer_or_null,
                "notice_period": integer_or_null,
                "priority": true_or_false_or_null
            }}
            """

            try:
                model = genai.GenerativeModel("models/gemini-2.5-flash-lite")
                response = model.generate_content(
                    extract_prompt,
                    generation_config=genai.types.GenerationConfig(
                        temperature=0.1,
                        max_output_tokens=200,
                        response_mime_type="application/json",
                    ),
                )

                try:
                    result_text = response.text  # may raise if no valid Part
                except Exception as inner_e:
                    print("JD save response.text access failed:", inner_e)
                    result_text = None

                if result_text:
                    result_text = result_text.strip()
                    clean_text = re.sub(r"```json|```", "", result_text).strip()
                    extracted = json.loads(clean_text)

                    # Merge extracted values into fields if missing
                    for key in ["name", "experience", "technology", "No_of_openings", "notice_period", "priority"]:
                        if not fields.get(key) and extracted.get(key) is not None:
                            fields[key] = extracted.get(key)

                    experience = fields.get('experience')
                    technology = fields.get('technology')
            except Exception as e:
                # If extraction fails, we just proceed with whatever we have
                print("JD save extraction failed:", e)

        # Simple heuristic fallback if LLM did not provide values
        if (not experience or not technology) and jd_text:
            lowered = jd_text.lower()
            # Try to infer technology from a small set of common stacks
            tech_map = ["python", "java", "react", "node", ".net", "aws"]
            if not technology:
                technology = next((t.title() for t in tech_map if t in lowered), technology)

            # Try to infer experience strings like "4+ years" or "3-5 years"
            if not experience:
                match = re.search(r"(\d+\+?\s*-?\s*\d*\s*years?)", jd_text, re.IGNORECASE)
                if match:
                    experience = match.group(1).strip()

        # After best-effort extraction, still require experience and technology
        if not experience or not technology:
            return Response(
                {
                    "error": "Missing required fields",
                    "details": "Could not infer 'experience' and 'technology' from JD text.",
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Auto-suggest JD name if not provided
        if not fields.get('name') and technology:
            fields["name"] = f"{technology} Developer"

        try:
            # Ensure technology is short and focused (fits into CharField(max_length=50))
            if isinstance(technology, str):
                # Keep at most 3 comma-separated items and trim each
                parts = [p.strip() for p in technology.split(',') if p.strip()]
                parts = parts[:3]
                # Limit each part length a bit to avoid one very long label
                parts = [p[:20] for p in parts]
                technology = ", ".join(parts)[:50]

            requirement = Requirement.objects.create(
                name=fields.get('name', ''),
                experience=experience or '',
                technology=technology or '',
                No_of_openings=fields.get('No_of_openings'),
                notice_period=fields.get('notice_period'),
                priority=fields.get('priority', False),
                base_text=jd_text  # Store final JD text (including user edits)
            )
            
            serializer = RequirementSerializer(requirement)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
            
        except Exception as e:
            return Response({"error": f"Save failed: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ChatAiView(APIView):
    def post(self, request):
        print('---------------------')
        serializer = ChatAiSerializer(data=request.data)
        if serializer.is_valid():
            user_message = serializer.validated_data['question']

            # Generate response using the model
            model = genai.GenerativeModel("models/gemini-2.5-flash")
            chat = model.start_chat()
            response = chat.send_message(user_message)

            data = {
                "question": user_message,
                "response": response.text,
            }

            return Response({"data": data}, status=status.HTTP_200_OK)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)